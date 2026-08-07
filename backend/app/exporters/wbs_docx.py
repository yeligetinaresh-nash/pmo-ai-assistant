from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _normalize_text(value) -> str:
    if value is None:
        return ""

    return str(value).strip()


def _format_effort(value) -> str:
    if value is None:
        return ""

    try:
        number = float(value)

        if number.is_integer():
            return str(int(number))

        return str(round(number, 2))

    except (TypeError, ValueError):
        return _normalize_text(value)


def _add_bullet_list(
    document: Document,
    items: list,
) -> None:
    if not items:
        document.add_paragraph("Not specified")
        return

    for item in items:
        paragraph = document.add_paragraph(
            style="List Bullet"
        )

        run = paragraph.add_run(
            _normalize_text(item)
        )

        run.font.name = "Arial"
        run.font.size = Pt(9)


def _set_cell_text(
    cell,
    value,
    bold: bool = False,
    font_size: float = 8,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

    run = paragraph.add_run(
        _normalize_text(value)
    )

    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def _shade_cell(
    cell,
    fill: str,
) -> None:
    cell_properties = (
        cell._tc.get_or_add_tcPr()
    )

    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        fill,
    )

    cell_properties.append(shading)


def _repeat_table_header(row) -> None:
    row_properties = (
        row._tr.get_or_add_trPr()
    )

    table_header = OxmlElement(
        "w:tblHeader"
    )

    table_header.set(
        qn("w:val"),
        "true",
    )

    row_properties.append(
        table_header
    )


def _prevent_row_split(row) -> None:
    row_properties = (
        row._tr.get_or_add_trPr()
    )

    cannot_split = OxmlElement(
        "w:cantSplit"
    )

    cannot_split.set(
        qn("w:val"),
        "true",
    )

    row_properties.append(
        cannot_split
    )


def _set_repeat_header_and_no_split(
    table,
) -> None:
    if not table.rows:
        return

    _repeat_table_header(
        table.rows[0]
    )

    for row in table.rows:
        _prevent_row_split(row)


def _set_cell_width(
    cell,
    width_inches: float,
) -> None:
    cell.width = Inches(
        width_inches
    )

    cell_properties = (
        cell._tc.get_or_add_tcPr()
    )

    cell_width = cell_properties.find(
        qn("w:tcW")
    )

    if cell_width is None:
        cell_width = OxmlElement(
            "w:tcW"
        )

        cell_properties.append(
            cell_width
        )

    cell_width.set(
        qn("w:w"),
        str(
            int(
                width_inches * 1440
            )
        ),
    )

    cell_width.set(
        qn("w:type"),
        "dxa",
    )


def _set_keep_with_next(
    paragraph,
) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _add_document_heading(
    document: Document,
    text: str,
    level: int = 1,
):
    heading = document.add_heading(
        text,
        level=level,
    )

    _set_keep_with_next(
        heading
    )

    return heading


def _calculate_remaining_effort(
    items: list[dict],
) -> float:
    parent_ids = {
        item.get("parent_wbs_id")
        for item in items
        if item.get("parent_wbs_id")
    }

    leaf_items = [
        item
        for item in items
        if item.get("wbs_id")
        not in parent_ids
    ]

    remaining_items = [
        item
        for item in leaf_items
        if item.get("status")
        in {
            "Planned",
            "In Progress",
            "On Hold",
        }
    ]

    total = sum(
        float(
            item.get(
                "estimated_effort_hours",
                0,
            )
            or 0
        )
        for item in remaining_items
    )

    return round(
        total,
        2,
    )


def _write_multiline_cell(
    cell,
    values: list,
    font_size: float = 7.5,
) -> None:
    cell.text = ""

    if not values:
        return

    for index, value in enumerate(values):
        if index == 0:
            paragraph = cell.paragraphs[0]
        else:
            paragraph = cell.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)

        run = paragraph.add_run(
            f"• {_normalize_text(value)}"
        )

        run.font.name = "Arial"
        run.font.size = Pt(font_size)

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.TOP
    )


def generate_wbs_docx(
    wbs: dict,
    output_path: str | Path,
) -> Path:
    output_path = Path(output_path)

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    section = document.sections[0]
    section.orientation = (
        WD_ORIENT.LANDSCAPE
    )

    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    for style_name in [
        "Heading 1",
        "Heading 2",
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"

    title = document.add_paragraph()
    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    title.paragraph_format.space_after = Pt(2)

    title_run = title.add_run(
        _normalize_text(
            wbs.get(
                "project_title",
                "Work Breakdown Structure",
            )
        )
    )

    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(17)

    subtitle = document.add_paragraph()
    subtitle.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle.paragraph_format.space_after = Pt(8)

    subtitle_run = subtitle.add_run(
        "WORK BREAKDOWN STRUCTURE"
    )

    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)

    items = list(
        wbs.get(
            "items",
            [],
        )
    )

    remaining_effort = (
        _calculate_remaining_effort(
            items
        )
    )

    summary_table = document.add_table(
        rows=5,
        cols=2,
    )

    summary_table.style = "Table Grid"
    summary_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    summary_table.autofit = False

    summary_data = [
        (
            "Artifact Status",
            wbs.get(
                "artifact_status"
            )
            or "Draft",
        ),
        (
            "Decomposition Approach",
            wbs.get(
                "decomposition_approach"
            )
            or "Deliverable-based",
        ),
        (
            "Full-Scope Estimated Effort",
            (
                f"{_format_effort(
                    wbs.get(
                        'total_estimated_effort_hours'
                    )
                )} hours"
            ),
        ),
        (
            "Remaining Estimated Effort",
            (
                f"{_format_effort(
                    remaining_effort
                )} hours"
            ),
        ),
        (
            "Purpose",
            wbs.get(
                "wbs_purpose"
            )
            or "Not specified",
        ),
    ]

    for row, values in zip(
        summary_table.rows,
        summary_data,
    ):
        _set_cell_width(
            row.cells[0],
            2.15,
        )

        _set_cell_width(
            row.cells[1],
            8.5,
        )

        _set_cell_text(
            row.cells[0],
            values[0],
            bold=True,
            font_size=8.5,
        )

        _shade_cell(
            row.cells[0],
            "D9EAF7",
        )

        _set_cell_text(
            row.cells[1],
            values[1],
            font_size=8.5,
        )

        _prevent_row_split(
            row
        )

    document.add_paragraph()

    _add_document_heading(
        document,
        "1. Work Breakdown Structure",
        level=1,
    )

    if not items:
        document.add_paragraph(
            "No WBS items available."
        )

    else:
        table = document.add_table(
            rows=1,
            cols=7,
        )

        table.style = "Table Grid"
        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        table.autofit = False

        headers = [
            "WBS ID",
            "Work Item",
            "Type",
            "Status",
            "Owner",
            "Effort",
            "Description",
        ]

        widths = [
            0.65,
            1.8,
            0.85,
            0.8,
            1.1,
            0.6,
            4.75,
        ]

        header_cells = (
            table.rows[0].cells
        )

        for index, header in enumerate(
            headers
        ):
            _set_cell_width(
                header_cells[index],
                widths[index],
            )

            _set_cell_text(
                header_cells[index],
                header,
                bold=True,
                font_size=8,
                alignment=(
                    WD_ALIGN_PARAGRAPH.CENTER
                ),
            )

            _shade_cell(
                header_cells[index],
                "B4C6E7",
            )

        _repeat_table_header(
            table.rows[0]
        )

        _prevent_row_split(
            table.rows[0]
        )

        for item in items:
            row = table.add_row()
            row_cells = row.cells

            row_values = [
                item.get(
                    "wbs_id",
                    "",
                ),
                item.get(
                    "name",
                    "",
                ),
                item.get(
                    "item_type",
                    "",
                ),
                item.get(
                    "status",
                    "",
                ),
                item.get(
                    "owner",
                    "",
                ),
                _format_effort(
                    item.get(
                        "estimated_effort_hours"
                    )
                ),
                item.get(
                    "description",
                    "",
                ),
            ]

            for index, value in enumerate(
                row_values
            ):
                _set_cell_width(
                    row_cells[index],
                    widths[index],
                )

                alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if index
                    in {
                        0,
                        2,
                        3,
                        5,
                    }
                    else WD_ALIGN_PARAGRAPH.LEFT
                )

                _set_cell_text(
                    row_cells[index],
                    value,
                    font_size=7.5,
                    alignment=alignment,
                )

            level = item.get(
                "level",
                3,
            )

            if level == 1:
                for cell in row_cells:
                    _shade_cell(
                        cell,
                        "D9EAD3",
                    )

                    for paragraph in (
                        cell.paragraphs
                    ):
                        for run in paragraph.runs:
                            run.bold = True

            elif level == 2:
                for cell in row_cells:
                    _shade_cell(
                        cell,
                        "EAF2F8",
                    )

            _prevent_row_split(
                row
            )

    document.add_page_break()

    _add_document_heading(
        document,
        "2. Work Package Details",
        level=1,
    )

    details_table = document.add_table(
        rows=1,
        cols=4,
    )

    details_table.style = "Table Grid"
    details_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    details_table.autofit = False

    detail_headers = [
        "WBS ID",
        "Work Item",
        "Acceptance Criteria",
        "Dependencies",
    ]

    detail_widths = [
        0.7,
        2.25,
        5.35,
        2.0,
    ]

    detail_header_cells = (
        details_table.rows[0].cells
    )

    for index, header in enumerate(
        detail_headers
    ):
        _set_cell_width(
            detail_header_cells[index],
            detail_widths[index],
        )

        _set_cell_text(
            detail_header_cells[index],
            header,
            bold=True,
            font_size=8,
            alignment=(
                WD_ALIGN_PARAGRAPH.CENTER
            ),
        )

        _shade_cell(
            detail_header_cells[index],
            "B4C6E7",
        )

    _repeat_table_header(
        details_table.rows[0]
    )

    _prevent_row_split(
        details_table.rows[0]
    )

    for item in items:
        criteria = item.get(
            "acceptance_criteria",
            [],
        )

        dependencies = item.get(
            "dependencies",
            [],
        )

        if not criteria and not dependencies:
            continue

        row = details_table.add_row()
        cells = row.cells

        for index, width in enumerate(
            detail_widths
        ):
            _set_cell_width(
                cells[index],
                width,
            )

        _set_cell_text(
            cells[0],
            item.get(
                "wbs_id",
                "",
            ),
            font_size=7.5,
            alignment=(
                WD_ALIGN_PARAGRAPH.CENTER
            ),
        )

        _set_cell_text(
            cells[1],
            item.get(
                "name",
                "",
            ),
            font_size=7.5,
        )

        _write_multiline_cell(
            cells[2],
            criteria,
            font_size=7.5,
        )

        _write_multiline_cell(
            cells[3],
            dependencies,
            font_size=7.5,
        )

        level = item.get(
            "level",
            3,
        )

        if level == 1:
            for cell in cells:
                _shade_cell(
                    cell,
                    "D9EAD3",
                )

        elif level == 2:
            for cell in cells:
                _shade_cell(
                    cell,
                    "EAF2F8",
                )

        _prevent_row_split(
            row
        )

    document.add_page_break()

    _add_document_heading(
        document,
        "3. Assumptions",
        level=1,
    )

    _add_bullet_list(
        document,
        wbs.get(
            "assumptions",
            [],
        ),
    )

    _add_document_heading(
        document,
        "4. Constraints",
        level=1,
    )

    _add_bullet_list(
        document,
        wbs.get(
            "constraints",
            [],
        ),
    )

    _add_document_heading(
        document,
        "5. Notes",
        level=1,
    )

    _add_bullet_list(
        document,
        wbs.get(
            "notes",
            [],
        ),
    )

    document.save(
        output_path
    )

    return output_path