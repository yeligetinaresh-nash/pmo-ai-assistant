from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _normalize_text(value: str) -> str:
    if not value:
        return ""

    replacements = {
        "chunkes text": "chunks text",
    }

    result = value

    for old_text, new_text in replacements.items():
        result = result.replace(
            old_text,
            new_text,
        )

    return result


def _add_bullet_list(
    document: Document,
    items: list[str],
) -> None:
    if not items:
        document.add_paragraph(
            "Not specified"
        )
        return

    for item in items:
        normalized_item = _normalize_text(
            str(item)
        )

        document.add_paragraph(
            normalized_item,
            style="List Bullet",
        )


def generate_project_charter_docx(
    charter: dict,
    output_path: str | Path,
) -> Path:
    output_path = Path(output_path)

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(
        _normalize_text(
            charter.get(
                "project_title",
                "Project Charter",
            )
        )
    )
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "PROJECT CHARTER"
    )
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)

    document.add_paragraph()

    summary_table = document.add_table(
        rows=4,
        cols=2,
    )
    summary_table.style = "Table Grid"

    summary_data = [
        (
            "Project Manager",
            _normalize_text(
                charter.get("project_manager")
                or "Not specified"
            ),
        ),
        (
            "Sponsor",
            _normalize_text(
                charter.get("sponsor")
                or "Not specified"
            ),
        ),
        (
            "Approval Status",
            _normalize_text(
                charter.get("approval_status")
                or "Draft"
            ),
        ),
        (
            "Artifact Status",
            "Draft",
        ),
    ]

    for row, values in zip(
        summary_table.rows,
        summary_data,
    ):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]

    document.add_heading(
        "1. Project Purpose",
        level=1,
    )

    project_purpose = _normalize_text(
        charter.get("project_purpose")
        or "Not specified"
    )

    document.add_paragraph(
        project_purpose
    )

    document.add_heading(
        "2. Project Summary",
        level=1,
    )

    project_summary = _normalize_text(
        charter.get("project_summary")
        or "Not specified"
    )

    document.add_paragraph(
        project_summary
    )

    document.add_heading(
        "3. Business Objectives",
        level=1,
    )

    _add_bullet_list(
        document,
        charter.get(
            "business_objectives",
            [],
        ),
    )

    document.add_heading(
        "4. In Scope",
        level=1,
    )

    _add_bullet_list(
        document,
        charter.get(
            "in_scope",
            [],
        ),
    )

    document.add_heading(
        "5. Out of Scope",
        level=1,
    )

    _add_bullet_list(
        document,
        charter.get(
            "out_of_scope",
            [],
        ),
    )

    document.add_heading(
        "6. Key Deliverables",
        level=1,
    )

    _add_bullet_list(
        document,
        charter.get(
            "key_deliverables",
            [],
        ),
    )

    document.add_heading(
        "7. Stakeholders",
        level=1,
    )

    stakeholders = charter.get(
        "stakeholders",
        [],
    )

    if stakeholders:
        stakeholder_table = document.add_table(
            rows=1,
            cols=3,
        )
        stakeholder_table.style = "Table Grid"

        headers = stakeholder_table.rows[0].cells
        headers[0].text = "Name"
        headers[1].text = "Role"
        headers[2].text = "Responsibility"

        for stakeholder in stakeholders:
            cells = stakeholder_table.add_row().cells

            cells[0].text = _normalize_text(
                stakeholder.get(
                    "name",
                    "",
                )
            )

            cells[1].text = _normalize_text(
                stakeholder.get(
                    "role",
                    "",
                )
            )

            cells[2].text = _normalize_text(
                stakeholder.get(
                    "responsibility",
                    "",
                )
            )
    else:
        document.add_paragraph(
            "Not specified"
        )

    document.add_heading(
        "8. High-Level Milestones",
        level=1,
    )

    milestones = charter.get(
        "milestones",
        [],
    )

    if milestones:
        milestone_table = document.add_table(
            rows=1,
            cols=3,
        )
        milestone_table.style = "Table Grid"

        headers = milestone_table.rows[0].cells
        headers[0].text = "Milestone"
        headers[1].text = "Description"
        headers[2].text = "Target Timeline"

        for milestone in milestones:
            cells = milestone_table.add_row().cells

            cells[0].text = _normalize_text(
                milestone.get(
                    "name",
                    "",
                )
            )

            cells[1].text = _normalize_text(
                milestone.get(
                    "description",
                    "",
                )
            )

            cells[2].text = _normalize_text(
                milestone.get(
                    "target_timeline",
                )
                or "To be confirmed"
            )
    else:
        document.add_paragraph(
            "Not specified"
        )

    sections = [
        (
            "9. Assumptions",
            "assumptions",
        ),
        (
            "10. Constraints",
            "constraints",
        ),
        (
            "11. Dependencies",
            "dependencies",
        ),
        (
            "12. High-Level Risks",
            "high_level_risks",
        ),
        (
            "13. Success Criteria",
            "success_criteria",
        ),
    ]

    for heading, key in sections:
        document.add_heading(
            heading,
            level=1,
        )

        _add_bullet_list(
            document,
            charter.get(
                key,
                [],
            ),
        )

    document.add_heading(
        "14. Approval",
        level=1,
    )

    approval_table = document.add_table(
        rows=3,
        cols=3,
    )
    approval_table.style = "Table Grid"

    approval_table.rows[0].cells[0].text = "Role"
    approval_table.rows[0].cells[1].text = "Name"
    approval_table.rows[0].cells[2].text = "Date / Signature"

    approval_table.rows[1].cells[0].text = "Project Manager"

    approval_table.rows[1].cells[1].text = _normalize_text(
        charter.get("project_manager")
        or "Not specified"
    )

    approval_table.rows[2].cells[0].text = "Sponsor"

    approval_table.rows[2].cells[1].text = _normalize_text(
        charter.get("sponsor")
        or "Not specified"
    )

    document.save(
        output_path
    )

    return output_path