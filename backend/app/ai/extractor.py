import fitz
from docx import Document


def extract_pdf_text(file_path: str) -> str:
    document = fitz.open(file_path)
    text_parts: list[str] = []

    try:
        for page in document:
            page_text = page.get_text()

            if page_text:
                text_parts.append(page_text)
    finally:
        document.close()

    return "\n".join(text_parts)


def extract_docx_text(file_path: str) -> str:
    document = Document(file_path)
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                text_parts.append(" | ".join(cells))

    return "\n".join(text_parts)